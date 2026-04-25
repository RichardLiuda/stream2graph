from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _set_run_font(run, *, size_pt: int = 11, name: str = "宋体"):
    run.font.name = name
    run.font.size = Pt(size_pt)


def main() -> None:
    out_path = Path("AI工具使用说明_填好版.docx").resolve()

    doc = Document()

    title = doc.add_paragraph("中国大学生计算机设计大赛")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.runs[0], size_pt=14)

    title2 = doc.add_paragraph("AI工具使用说明(2026年版)")
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title2.runs[0], size_pt=14)

    meta = doc.add_paragraph("作品编号：        作品名称：")
    _set_run_font(meta.runs[0], size_pt=11)

    doc.add_paragraph(
        "说明：本项目使用 AI 工具主要用于代码定位、方案对比、生成少量样板代码与交互/样式优化建议，以提升开发效率。"
        "需求定义、交互设计取舍、视觉风格把控、参数标定、功能验收与最终责任均由团队成员完成。"
        "AI 输出在纳入代码前均经过人工审查、修改与测试。"
    )

    headers = [
        "序号",
        "AI工具的名称、版本、访问方式（网页、API或客户端），使用时间",
        "使用AI工具的环节与目的（立项构思、文献综述、语言润色、内容生成、图表优化、代码编程、数据分析等）",
        "关键提示词",
        "AI回复的关键内容（在此简要说明，并在附录中给出佐证）",
        "AI回复的人工修改说明",
        "采纳比例与说明",
    ]

    rows = [
        [
            "1",
            "IDE 智能编程辅助功能（本地客户端），2026年3月下旬（约3/20–3/31）",
            "立项开发：梳理平台模块与路由结构，明确“实时工作台/样本对比/研究页/导出”页面边界",
            "“梳理正式平台的页面入口与模块功能边界，并给出实现注意点”",
            "输出模块划分建议（首页、登录、实时工作台、样本对比、报告导出、研究参与者页）与实现注意点",
            "团队结合实际需求删改模块边界与页面内容，最终以人工产品设计为准",
            "15%：主要采纳结构梳理思路；具体信息架构由人工确定",
        ],
        [
            "2",
            "IDE 智能编程辅助功能（本地客户端），2026年4月上旬（约4/1–4/6）",
            "代码编程：实时工作台 UI 迭代（按钮、布局、交互反馈），提升可用性与演示稳定性",
            "“实时工作台布局怎么优化：左右布局比例、按钮层级、交互提示更清晰”",
            "给出布局优化方案与可落地的组件结构/样式建议",
            "团队按实际屏幕尺寸与交互路径多次调整间距与位置，保留最符合演示场景的方案",
            "20%：AI 提供备选方案；最终 UI 由人工多轮验收决定",
        ],
        [
            "3",
            "IDE 智能编程辅助功能（本地客户端），2026年4月中旬（约4/15–4/17）",
            "代码编程：批注/注释能力接入（工具选择、笔宽/颜色、橡皮等），保证基本可用",
            "“批注工具（笔/框/文字/橡皮）怎么组织交互与状态管理”",
            "输出批注工具组织方式（工具状态、参数面板、撤销/清空）与实现建议",
            "团队根据“始终可用/不依赖会话”等要求调整逻辑，最终以人工需求为准",
            "25%：AI 提供实现思路；关键交互规则由人工定义并验收",
        ],
        [
            "4",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/19–4/23）",
            "图表优化/界面优化：首页流程演示动效组件（卡片、箭头、分支小卡）增强观感与可读性",
            "“小卡不要遮挡：上面那张下移、两张右移；再微调若干档位”",
            "定位布局参数并提供可调偏移与间距，支持多轮快速微调",
            "团队基于实际视觉效果反复调整档位与像素，最终值由人工确定",
            "25%：AI 负责定位与样板修改；最终参数由人工精调",
        ],
        [
            "5",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/22–4/25）",
            "图表优化：分支箭头锚点对齐与越界修正（提升“语音”分支指向清晰度）",
            "“语音箭头起点贴主卡右上角、不要出界；更贴一点/收回一点”",
            "给出基于卡片边界计算锚点的实现方式，并支持像素级回调",
            "团队通过截图比对校准锚点，保证既贴角又不越界",
            "20%：AI 给出实现方式；最终像素对齐由人工校准",
        ],
        [
            "6",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/22–4/25）",
            "代码编程：临时关闭“小卡拖拽位置记忆”，避免刷新后布局漂移影响演示",
            "“先不要记忆我手动调整的功能”",
            "找到持久化存储点并关闭读写，使布局回到默认参数",
            "团队确认是阶段性需求（演示期优先稳定），后续可再启用",
            "30%：AI 完成实现；人工确认需求与验收",
        ],
        [
            "7",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/20–4/25）",
            "界面优化：批注按钮分组与莫兰迪分色（删除类偏红），降低“按钮杂乱”",
            "“按键太多太杂；用底卡抱团；每个按钮按功能分色，删除偏红；撤销清空横排”",
            "提供分组容器样式与按功能分色方案，并修复按钮文字竖排问题",
            "团队按美观要求反复调整对比度、边框深浅与色彩强度，确保不刺眼且可区分",
            "35%：AI 给出结构+初稿；最终色彩与强度由人工决定",
        ],
        [
            "8",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/24–4/25）",
            "交互优化：按钮保持原形，点击后下方弹出紧凑卡片放调节项，避免挤压布局",
            "“保留原按钮形状；点击后下面弹出卡片，放调节项，卡片紧凑”",
            "给出“固定按钮 + 下方面板承载调节项”的结构建议（笔宽/颜色、框宽/颜色、文字颜色、橡皮预设）",
            "团队控制卡片尺寸与样式细节，确保紧凑且不遮挡关键区域",
            "30%：AI 提供结构建议；人工完成最终观感与尺寸验收",
        ],
        [
            "9",
            "IDE 智能编程辅助功能（本地客户端），2026年4月下旬（约4/20–4/25）",
            "体验清理：删除“正常状态但像报错”的提示（如空 Mermaid 时误导性提示），避免用户误判",
            "“把没有 Mermaid 代码、渲染已就绪之类正常提示删掉/降级”",
            "定位空态/状态徽章并建议仅保留真实失败提示",
            "团队确认“空内容=正常空态”，保留“真实编译/渲染失败”的提示并做回归检查",
            "25%：AI 帮助定位与建议；最终口径与验收由人工负责",
        ],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, size_pt=10)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i != 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            _set_run_font(run, size_pt=10)

    doc.add_paragraph("")
    doc.add_paragraph("附录2：")
    doc.add_paragraph("序号1的佐证材料：对话截图 + 关键 diff 截图（页面结构/路由/模块梳理）")
    doc.add_paragraph("序号2的佐证材料：实时工作台布局迭代对话截图 + 关键 diff 截图")
    doc.add_paragraph("序号3的佐证材料：批注功能接入对话截图 + 关键 diff 截图")
    doc.add_paragraph("序号4的佐证材料：小卡遮挡问题前/后截图 + 对话调参截图 + diff")
    doc.add_paragraph("序号5的佐证材料：箭头锚点对齐前/后截图 + 对话截图 + diff")
    doc.add_paragraph("序号6的佐证材料：关闭位置记忆的对话截图 + diff")
    doc.add_paragraph("序号7的佐证材料：批注工具栏分组/分色截图 + 对话截图 + diff")
    doc.add_paragraph("序号8的佐证材料：点击下方弹出设置卡片截图 + 对话截图 + diff")
    doc.add_paragraph("序号9的佐证材料：空 Mermaid 时不再出现误导提示的截图 + diff")

    foot = doc.add_paragraph(f"生成日期：{date.today().isoformat()}")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(foot.runs[0], size_pt=9)

    doc.save(out_path)
    print(f"wrote: {out_path}")


if __name__ == "__main__":
    main()

